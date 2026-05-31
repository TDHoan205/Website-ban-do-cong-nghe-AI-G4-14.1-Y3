from docx import Document
from docx.shared import Pt

def set_normal_style(doc: Document):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(13)


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(text)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for r in rows:
        row_cells = table.add_row().cells
        for i, v in enumerate(r):
            row_cells[i].text = v

    doc.add_paragraph("")


def main():
    doc = Document()
    set_normal_style(doc)

    add_heading(doc, "CHƯƠNG II: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG", 1)

    add_heading(doc, "2.2. Mô tả cấu trúc dữ liệu của chương trình", 2)
    add_paragraph(
        doc,
        "Hệ thống sử dụng Microsoft SQL Server với cơ sở dữ liệu TechShopWebsite2. "
        "CSDL được thiết kế theo mô hình quan hệ để phục vụ các nghiệp vụ: quản lý tài khoản, sản phẩm, "
        "giỏ hàng, đơn hàng, thanh toán và dữ liệu phục vụ AI Chatbot (RAG). "
        "Cấu trúc CSDL được định nghĩa trong file SQL/schema.sql và dữ liệu mẫu trong SQL/seed_data.sql.",
    )

    add_heading(doc, "2.2.2. Các nhóm bảng dữ liệu chính", 3)
    add_paragraph(doc, "(1) Nhóm người dùng – phân quyền")

    add_heading(doc, "Bảng Roles", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["role_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["role_name", "NVARCHAR(50)", "NOT NULL, UNIQUE"],
            ["description", "NVARCHAR(255)", "NULL"],
        ],
    )

    add_heading(doc, "Bảng Accounts", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["account_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["username", "NVARCHAR(50)", "NOT NULL, UNIQUE"],
            ["password_hash", "NVARCHAR(255)", "NOT NULL"],
            ["email", "NVARCHAR(100)", "NULL, UNIQUE"],
            ["full_name", "NVARCHAR(100)", "NULL"],
            ["phone", "NVARCHAR(20)", "NULL"],
            ["address", "NVARCHAR(255)", "NULL"],
            ["is_active", "BIT", "NOT NULL, DEFAULT 1"],
            ["role_id", "INT", "NOT NULL, FK → Roles(role_id)"],
            ["reset_token", "NVARCHAR(64)", "NULL"],
            ["reset_token_expiry", "DATETIME2", "NULL"],
            ["created_at", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
            ["updated_at", "DATETIME2", "NULL"],
        ],
    )

    add_heading(doc, "Bảng Employees", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["employee_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["account_id", "INT", "NOT NULL, UNIQUE, FK → Accounts(account_id)"],
            ["department", "NVARCHAR(50)", "NULL"],
            ["position", "NVARCHAR(50)", "NULL"],
            ["hire_date", "DATETIME2", "NULL"],
            ["salary", "INT", "NULL"],
            ["is_active", "BIT", "NOT NULL, DEFAULT 1"],
            ["created_at", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
        ],
    )

    add_heading(doc, "Bảng Users", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["user_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["username", "NVARCHAR(50)", "NOT NULL, UNIQUE"],
            ["email", "NVARCHAR(100)", "NOT NULL, UNIQUE"],
            ["password_hash", "NVARCHAR(255)", "NOT NULL"],
            ["full_name", "NVARCHAR(100)", "NULL"],
            ["phone", "NVARCHAR(20)", "NULL"],
            ["address", "NVARCHAR(255)", "NULL"],
            ["is_active", "BIT", "NOT NULL, DEFAULT 1"],
            ["role", "NVARCHAR(20)", "NOT NULL, DEFAULT 'Customer'"],
            ["avatar_url", "NVARCHAR(500)", "NULL"],
            ["reset_token", "NVARCHAR(64)", "NULL"],
            ["reset_token_expiry", "DATETIME2", "NULL"],
            ["created_at", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
            ["updated_at", "DATETIME2", "NULL"],
        ],
    )

    add_paragraph(doc, "(2) Nhóm danh mục – nhà cung cấp – sản phẩm")

    add_heading(doc, "Bảng Categories", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["category_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["name", "NVARCHAR(100)", "NOT NULL, UNIQUE"],
            ["description", "NVARCHAR(MAX)", "NULL"],
            ["image_url", "NVARCHAR(500)", "NULL"],
            ["display_order", "INT", "NOT NULL, DEFAULT 0"],
            ["is_active", "BIT", "NOT NULL, DEFAULT 1"],
            ["created_at", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
            ["updated_at", "DATETIME2", "NULL"],
        ],
    )

    add_heading(doc, "Bảng Suppliers", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["supplier_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["name", "NVARCHAR(255)", "NOT NULL"],
            ["contact_person", "NVARCHAR(100)", "NULL"],
            ["phone", "NVARCHAR(20)", "NULL"],
            ["email", "NVARCHAR(100)", "NULL"],
            ["address", "NVARCHAR(255)", "NULL"],
            ["is_active", "BIT", "NOT NULL, DEFAULT 1"],
            ["created_at", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
            ["updated_at", "DATETIME2", "NULL"],
        ],
    )

    add_heading(doc, "Bảng Products", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["product_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["name", "NVARCHAR(255)", "NOT NULL"],
            ["description", "NVARCHAR(MAX)", "NULL"],
            ["image_url", "NVARCHAR(500)", "NULL"],
            ["price", "DECIMAL(18,2)", "NOT NULL"],
            ["original_price", "DECIMAL(18,2)", "NULL"],
            ["stock_quantity", "INT", "NOT NULL, DEFAULT 0"],
            ["is_available", "BIT", "NOT NULL, DEFAULT 1"],
            ["rating", "DECIMAL(3,1)", "NOT NULL, DEFAULT 4.5"],
            ["is_new", "BIT", "NOT NULL, DEFAULT 0"],
            ["is_hot", "BIT", "NOT NULL, DEFAULT 0"],
            ["discount_percent", "INT", "NOT NULL, DEFAULT 0"],
            ["specifications", "NVARCHAR(MAX)", "NULL"],
            ["category_id", "INT", "NULL, FK → Categories(category_id)"],
            ["supplier_id", "INT", "NULL, FK → Suppliers(supplier_id)"],
            ["created_at", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
            ["updated_at", "DATETIME2", "NULL"],
        ],
    )

    add_heading(doc, "Bảng ProductVariants", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["variant_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["product_id", "INT", "NOT NULL, FK → Products(product_id), ON DELETE CASCADE"],
            ["color", "NVARCHAR(50)", "NULL"],
            ["color_hex", "NVARCHAR(7)", "NULL"],
            ["storage", "NVARCHAR(20)", "NULL"],
            ["ram", "NVARCHAR(20)", "NULL"],
            ["variant_name", "NVARCHAR(100)", "NULL"],
            ["sku", "NVARCHAR(50)", "NULL"],
            ["price", "DECIMAL(18,2)", "NULL"],
            ["original_price", "DECIMAL(18,2)", "NULL"],
            ["stock_quantity", "INT", "NOT NULL, DEFAULT 0"],
            ["display_order", "INT", "NOT NULL, DEFAULT 0"],
            ["is_active", "BIT", "NOT NULL, DEFAULT 1"],
            ["created_at", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
        ],
    )

    add_heading(doc, "Bảng ProductImages", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["image_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["product_id", "INT", "NOT NULL, FK → Products(product_id), ON DELETE NO ACTION"],
            ["variant_id", "INT", "NULL, FK → ProductVariants(variant_id), ON DELETE SET NULL"],
            ["image_url", "NVARCHAR(500)", "NOT NULL"],
            ["display_order", "INT", "NOT NULL, DEFAULT 0"],
            ["is_primary", "BIT", "NOT NULL, DEFAULT 0"],
            ["created_at", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
        ],
    )

    add_paragraph(doc, "(3) Nhóm kho hàng")

    add_heading(doc, "Bảng Inventory", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["inventory_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["product_id", "INT", "NOT NULL, UNIQUE, FK → Products(product_id)"],
            ["quantity_in_stock", "INT", "NOT NULL, DEFAULT 0"],
            ["min_stock_level", "INT", "NOT NULL, DEFAULT 5"],
            ["max_stock_level", "INT", "NOT NULL, DEFAULT 100"],
            ["last_restock_date", "DATETIME2", "NULL"],
            ["created_at", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
            ["updated_at", "DATETIME2", "NULL"],
        ],
    )

    add_paragraph(doc, "(4) Nhóm giỏ hàng – đơn hàng")

    add_heading(doc, "Bảng Carts", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["cart_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["account_id", "INT", "NOT NULL, FK → Accounts(account_id)"],
            ["created_at", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
            ["updated_at", "DATETIME2", "NULL"],
        ],
    )

    add_heading(doc, "Bảng CartItems", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["cart_item_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["cart_id", "INT", "NOT NULL, FK → Carts(cart_id), ON DELETE CASCADE"],
            ["product_id", "INT", "NOT NULL, FK → Products(product_id)"],
            ["variant_id", "INT", "NULL, FK → ProductVariants(variant_id)"],
            ["quantity", "INT", "NOT NULL, DEFAULT 1"],
            ["added_date", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
        ],
    )

    add_heading(doc, "Bảng Orders", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["order_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["account_id", "INT", "NULL, FK → Accounts(account_id)"],
            ["order_date", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
            ["total_amount", "DECIMAL(10,2)", "NOT NULL"],
            ["status", "NVARCHAR(20)", "NOT NULL, DEFAULT 'Pending'"],
            ["customer_name", "NVARCHAR(100)", "NULL"],
            ["customer_phone", "NVARCHAR(20)", "NULL"],
            ["customer_address", "NVARCHAR(255)", "NULL"],
            ["notes", "NVARCHAR(500)", "NULL"],
            ["created_at", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
            ["updated_at", "DATETIME2", "NULL"],
        ],
    )

    add_heading(doc, "Bảng OrderItems", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["order_item_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["order_id", "INT", "NOT NULL, FK → Orders(order_id), ON DELETE CASCADE"],
            ["product_id", "INT", "NOT NULL, FK → Products(product_id)"],
            ["variant_id", "INT", "NULL, FK → ProductVariants(variant_id)"],
            ["product_name", "NVARCHAR(255)", "NOT NULL"],
            ["variant_name", "NVARCHAR(100)", "NULL"],
            ["quantity", "INT", "NOT NULL"],
            ["unit_price", "DECIMAL(10,2)", "NOT NULL"],
            ["subtotal", "DECIMAL(10,2)", "NOT NULL"],
        ],
    )

    add_paragraph(doc, "(5) Nhóm dữ liệu thanh toán")

    add_heading(doc, "Bảng Payments", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["payment_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["order_id", "NVARCHAR(50)", "NOT NULL, UNIQUE"],
            ["account_id", "INT", "NULL, FK → Accounts(account_id)"],
            ["amount", "BIGINT", "NOT NULL"],
            ["payment_method", "NVARCHAR(20)", "NOT NULL, DEFAULT 'QR_BANKING'"],
            ["transaction_code", "NVARCHAR(100)", "NULL"],
            ["transfer_content", "NVARCHAR(200)", "NULL"],
            ["status", "NVARCHAR(20)", "NOT NULL, DEFAULT 'PENDING'"],
            ["qr_data", "NVARCHAR(MAX)", "NULL"],
            ["qr_image_base64", "NVARCHAR(MAX)", "NULL"],
            ["bank_code", "NVARCHAR(20)", "NULL"],
            ["bank_account", "NVARCHAR(50)", "NULL"],
            ["bank_account_name", "NVARCHAR(200)", "NULL"],
            ["created_at", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
            ["paid_at", "DATETIME2", "NULL"],
            ["expires_at", "DATETIME2", "NULL"],
            ["verified_by", "NVARCHAR(50)", "NULL"],
            ["notes", "NVARCHAR(500)", "NULL"],
        ],
    )

    add_paragraph(doc, "(6) Nhóm dữ liệu chatbot AI (RAG)")

    add_heading(doc, "Bảng ChatSessions", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["session_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["session_uuid", "NVARCHAR(36)", "NOT NULL, UNIQUE"],
            ["account_id", "INT", "NULL, FK → Accounts(account_id)"],
            ["device_info", "NVARCHAR(255)", "NULL"],
            ["started_at", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
            ["ended_at", "DATETIME2", "NULL"],
            ["is_active", "BIT", "NOT NULL, DEFAULT 1"],
        ],
    )

    add_heading(doc, "Bảng ChatMessages", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["message_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["session_id", "INT", "NOT NULL, FK → ChatSessions(session_id), ON DELETE CASCADE"],
            ["sender_type", "NVARCHAR(20)", "NOT NULL"],
            ["message_content", "NVARCHAR(MAX)", "NOT NULL"],
            ["intent", "NVARCHAR(50)", "NULL"],
            ["confidence_score", "NVARCHAR(10)", "NULL"],
            ["is_product_recommendation", "BIT", "NOT NULL, DEFAULT 0"],
            ["recommended_product_ids", "NVARCHAR(255)", "NULL"],
            ["created_at", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
        ],
    )

    add_heading(doc, "Bảng AIConversationLogs", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["log_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["session_id", "INT", "NOT NULL, FK → ChatSessions(session_id)"],
            ["account_id", "INT", "NULL, FK → Accounts(account_id)"],
            ["user_message", "NVARCHAR(MAX)", "NULL"],
            ["bot_response", "NVARCHAR(MAX)", "NULL"],
            ["intent_detected", "NVARCHAR(50)", "NULL"],
            ["confidence_score", "NVARCHAR(10)", "NULL"],
            ["response_time_ms", "INT", "NULL"],
            ["is_escalated_to_staff", "BIT", "NOT NULL, DEFAULT 0"],
            ["created_at", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
        ],
    )

    add_heading(doc, "Bảng FAQs", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["faq_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["question", "NVARCHAR(500)", "NOT NULL"],
            ["answer", "NVARCHAR(MAX)", "NOT NULL"],
            ["category", "NVARCHAR(50)", "NULL"],
            ["display_order", "INT", "NOT NULL, DEFAULT 0"],
            ["is_active", "BIT", "NOT NULL, DEFAULT 1"],
            ["created_at", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
            ["updated_at", "DATETIME2", "NULL"],
        ],
    )

    add_heading(doc, "Bảng Notifications", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["notification_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["account_id", "INT", "NULL, FK → Accounts(account_id)"],
            ["title", "NVARCHAR(200)", "NOT NULL"],
            ["content", "NVARCHAR(MAX)", "NOT NULL"],
            ["notification_type", "NVARCHAR(50)", "NULL"],
            ["is_read", "BIT", "NOT NULL, DEFAULT 0"],
            ["created_at", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
        ],
    )

    add_heading(doc, "Bảng KnowledgeChunks", 4)
    add_table(
        doc,
        ["Tên cột", "Kiểu dữ liệu", "Ràng buộc"],
        [
            ["chunk_id", "INT", "IDENTITY(1,1), PRIMARY KEY"],
            ["content", "NVARCHAR(MAX)", "NOT NULL"],
            ["content_type", "NVARCHAR(50)", "NULL"],
            ["source_id", "INT", "NULL"],
            ["source_table", "NVARCHAR(50)", "NULL"],
            ["embedding_vector", "NVARCHAR(MAX)", "NULL"],
            ["metadata_json", "NVARCHAR(MAX)", "NULL"],
            ["created_at", "DATETIME2", "NOT NULL, DEFAULT SYSUTCDATETIME()"],
        ],
    )

    add_heading(doc, "2.3. Mô tả cấu trúc mã nguồn chương trình và giải thích ý nghĩa các file", 2)
    add_paragraph(
        doc,
        "Dự án được tổ chức theo mô hình MVC kết hợp Service Layer trên nền FastAPI. "
        "Controller nhận HTTP request, gọi Service xử lý nghiệp vụ, Service thao tác dữ liệu qua SQLAlchemy ORM, "
        "sau đó Controller trả về giao diện HTML (Jinja2) hoặc JSON. Static files (CSS/JS/Images) được quản lý trong wwwroot.",
    )

    add_heading(doc, "2.3.1. Cấu trúc thư mục", 3)
    add_paragraph(doc, "- app.py: Entry point, cấu hình FastAPI, templates, static, middleware, include router")
    add_paragraph(doc, "- Controllers/: Tầng xử lý HTTP, định nghĩa routes theo từng module chức năng")
    add_paragraph(doc, "- Services/: Tầng nghiệp vụ, truy vấn DB và xử lý logic")
    add_paragraph(doc, "- Models/: Định nghĩa entity ORM ánh xạ bảng trong SQL Server")
    add_paragraph(doc, "- Data/: Cấu hình kết nối DB (engine, session)")
    add_paragraph(doc, "- Views/: Template HTML Jinja2 (giao diện người dùng và admin)")
    add_paragraph(doc, "- wwwroot/: Tài nguyên tĩnh (css/js/images/uploads)")
    add_paragraph(doc, "- SQL/: Script tạo schema và seed dữ liệu")

    out_path = "TechStoreAI_Chuong2_CauTrucDuLieu_MaNguon.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
